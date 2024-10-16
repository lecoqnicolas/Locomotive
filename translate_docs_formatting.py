import argparse
import logging
import os
import time
import torch
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain_community.document_loaders import TextLoader 
from docx import Document
from locomotive_llm.documents_handling import DocxLoader
from locomotive_llm.load import load_config
from locomotive_llm.model import get_pipeline

def read_docx_with_langchain(file_path):
    """Reads a .docx file using the custom LangChain loader."""
    loader = DocxLoader(file_path)
    documents = loader.load()
    return "\n".join([doc.page_content for doc in documents])

def read_document_with_langchain(file_path):
    """Reads .docx or .txt file."""
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == ".docx":
        return read_docx_with_langchain(file_path)
    elif file_extension == ".txt":
        loader = TextLoader(file_path)
        documents = loader.load()
        return "\n".join([doc.page_content for doc in documents])
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
# Function to read the document while preserving the formatting
def read_docx_with_formatting(file_path):
    """Reads a .docx file and preserves the structure (tables, paragraphs, etc.)."""
    doc = Document(file_path)
    elements = []

    for para in doc.paragraphs:
        elements.append({'type': 'paragraph', 'content': para.text})
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        elements.append({'type': 'table', 'content': table_data})

    return elements

# Batch text processing
def batch_texts(texts, batch_size):
    """Divides texts into batches of batch_size."""
    for i in range(0, len(texts), batch_size):
        yield texts[i:i + batch_size]

def translate_document_with_formatting(pipeline, elements, batch_size, src_lang, tgt_lang):
    translated_elements = []

    # Split paragraphs and table rows for sequential processing
    for el in elements:
        if el['type'] == 'paragraph' and el['content'].strip():
            # Translate paragraph in batch if needed
            paragraphs = list(batch_texts([el['content']], batch_size))
            for batch in paragraphs:
                translated_paragraphs = pipeline.transform(batch, src_lang, tgt_lang)
                for text in translated_paragraphs:
                    translated_elements.append({'type': 'paragraph', 'content': text})

        elif el['type'] == 'table':
            # Translate each table row one by one
            translated_table = []
            for row in el['content']:
                translated_row = pipeline.transform(row, src_lang, tgt_lang)
                translated_table.append(translated_row)
            translated_elements.append({'type': 'table', 'content': translated_table})

    return translated_elements
def translate_document(pipeline, text, batch_size, src_lang, tgt_lang):
    lines = [line for line in text.split("\n") if line.strip()]
    translated_batches = []

    for idx, batch in enumerate(batch_texts(lines, batch_size)):
        translated_texts = pipeline.transform(batch, src_lang, tgt_lang)
        translated_batches.append("\n".join(translated_texts))
    
    return "\n".join(translated_batches)

def translate_text_batch(pipeline, batch, src_lang, tgt_lang):
    if batch:
        return pipeline.transform(batch, src_lang, tgt_lang)
    return []


def translate_table_batch(pipeline, table, src_lang, tgt_lang):
    translated_table = []
    for row in table:
        translated_row = pipeline.transform(row, src_lang, tgt_lang)
        translated_table.append(translated_row)
    return translated_table


def write_docx_with_formatting(translated_elements, output_path):
    doc = Document()
    
    for element in translated_elements:
        if element['type'] == 'paragraph':
            doc.add_paragraph(element['content'])
        elif element['type'] == 'table':
            table = doc.add_table(rows=len(element['content']), cols=len(element['content'][0]))
            for i, row in enumerate(element['content']):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text

    doc.save(output_path)

# Writing a translated .txt file
def write_txt(translated_text, output_path):
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(translated_text)

def main(params: argparse.Namespace) -> None:
    logging.basicConfig(level="DEBUG")
    config = load_config(params.config, params.reverse)
    pipeline_class = get_pipeline(config)
    pipeline = pipeline_class(model_id=config.llm_model,
                              device="cuda" if torch.cuda.is_available() and not params.cpu else "cpu",
                              prompt_file=config.prompt,
                              batch_size=config.batch_size,
                              output_parser=config.response_parsing_method)

    input_file = params.input_file
    
    # Read the document (docx or txt)
    try:
        if input_file.endswith(".docx"):
            elements = read_docx_with_formatting(input_file)
        else:
            text = read_document_with_langchain(input_file)
    except ValueError as e:
        logging.error(str(e))
        exit(1)

    # Start the translation process
    logging.info("Starting translation of the document.")
    total_start_time = time.time()
    
    if input_file.endswith(".docx"):
        translated_elements = translate_document_with_formatting(pipeline, elements, config.batch_size, config.src_name, config.tgt_name)
        total_end_time = time.time()
        output_file = params.output_file
        write_docx_with_formatting(translated_elements, output_file)
    else:
        translated_text = translate_document(pipeline, text, config.batch_size, config.src_name, config.tgt_name)
        total_end_time = time.time()
        output_file = params.output_file
        write_txt(translated_text, output_file)

    logging.info(f"Total translation time: {total_end_time - total_start_time:.2f} seconds.")
    logging.info(f"Translation completed. Output saved at {output_file}.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Translate a document using TowerLLM model')
    
    parser.add_argument('--config',
                        type=str,
                        default="config/config_en_fr.yml",
                        help='Path to model-config.yml. Default: %(default)s')
    
    parser.add_argument('--reverse',
                        action='store_true',
                        help='Reverse the source and target languages in the configuration and data sources. Default: %(default)s')
    
    parser.add_argument('--input_file',
                        type=str,
                        required=True,
                        help='Path to the input document file (.docx or .txt).')
    
    parser.add_argument('--output_file',
                        type=str,
                        required=True,
                        help='Path to the output translated document file (.docx or .txt).')
    
    parser.add_argument('--cpu',
                        action="store_true",
                        help='Force CPU use. Default: %(default)s')
    
    args = parser.parse_args()

    main(args)
