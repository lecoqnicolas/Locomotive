import argparse
import logging
import os
import torch
from pathlib import Path
import time
import mlflow
from locomotive_llm.load import load_config
from locomotive_llm.model import get_pipeline
from locomotive_llm.utils import log_dataclass

from docx import Document


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    return "\n".join([para.text for para in doc.paragraphs])

def write_docx(translated_text, output_path):
    doc = Document()
    for paragraph in translated_text.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(output_path)


def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def write_txt(translated_text, output_path):
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(translated_text)

def clean_text(text, phrase_to_remove):
    return text.replace(phrase_to_remove, '').strip()

def batch_texts(texts, batch_size):
    """Divides texts into batches of batch_size."""
    for i in range(0, len(texts), batch_size):
        yield texts[i:i + batch_size]

def translate_document(pipeline, text, batch_size, src_lang, tgt_lang):
    # Splitting the text into batches for translation
    lines = [line for line in text.split("\n") if line.strip()]
    batches = list(batch_texts(lines, batch_size))
    translated_batches = []
    
    for idx, batch in enumerate(batches):
        start_time = time.time()  # Start the timer for each batch
        translated_texts = pipeline.transform(batch, src_lang, tgt_lang)
        end_time = time.time()    # End the timer for each batch
        translated_batches.append("\n".join(translated_texts))
        
        logging.info(f"Batch {idx + 1}/{len(batches)} processed in {end_time - start_time:.2f} seconds.")
    
    return "\n".join(translated_batches)


def main(params: argparse.Namespace) -> None:
    logging.basicConfig(level="DEBUG")
    config = load_config(params.config, params.reverse)
    pipeline_class = get_pipeline(config)
    pipeline = pipeline_class(model_id=config.llm_model,
                              device="cuda" if torch.cuda.is_available() and not params.cpu else "cpu",
                              prompt_file=config.prompt,
                              batch_size=config.batch_size)

    input_file = params.input_file
    file_extension = os.path.splitext(input_file)[1].lower()
    
    if file_extension == ".docx":
        text = read_docx(input_file)
    elif file_extension == ".txt":
        text = read_txt(input_file)
    else:
        logging.error("Unsupported file format. Only .docx and .txt files are supported.")
        exit(1)

    # Translate the text
    logging.info("Starting translation of the document.")
    total_start_time = time.time()
    translated_text = translate_document(pipeline, text, config.batch_size, config.src_name, config.tgt_name)
    total_end_time = time.time()
    logging.info(f"Total translation time: {total_end_time - total_start_time:.2f} seconds.")
    output_file = params.output_file
    if file_extension == ".docx":
        write_docx(translated_text, output_file)
    elif file_extension == ".txt":
        write_txt(translated_text, output_file)

    logging.info(f"Translation completed. Output saved at {output_file}.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Translate a document using TowerLLM model')
    
    parser.add_argument('--config',
                        type=str,
                        default="config/config_en_fr.yml",
                        help='Path to model-config.yml. Default: %(default)s')
    
    parser.add_argument('--reverse',
                        action='store_true',
                        help='Reverse the source and target languages in the configuration and data sources. Default: %(default)s')
    
    parser.add_argument('--input_file',
                        type=str,
                        required=True,
                        help='Path to the input document file (.docx or .txt).')
    
    parser.add_argument('--output_file',
                        type=str,
                        required=True,
                        help='Path to the output translated document file (.docx or .txt).')
    
    parser.add_argument('--cpu',
                        action="store_true",
                        help='Force CPU use. Default: %(default)s')
    
    args = parser.parse_args()

    main(args)
