from docx import Document


class DocumentTemplate:
    def __init__(self, file_path: str):
        self._doc = Document(file_path)
        self._nb_para = None
        self._nb_table = None

    def get_content(self):
        elements = []
        for para in self._doc.paragraphs:
            elements.append({'type': 'paragraph', 'content': para.text})
        self._nb_para = len(elements)

        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    elements.append({'type': 'table', 'content': cell.text})
        return elements

    def map_translations(self, translations):
        element_count = 0
        for para in self._doc.paragraphs:
            if len(translations[element_count]) > 0:
                para.text = translations[element_count]
            element_count += 1
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(translations[element_count]) > 0:
                        cell.text = translations[element_count]
                    element_count += 1
    def save(self, file_path):
        self._doc.save(file_path)
