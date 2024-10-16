from docx import Document

from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document as LangchainDocument


class DocxLoader(BaseLoader):
    def __init__(self, file_path: str):
        self.file_path = file_path

    def lazy_load(self):
        doc = Document(self.file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        yield LangchainDocument(page_content=text, metadata={"source": self.file_path})

    def load(self) -> list[LangchainDocument]:
        """Loads the entire .docx file content at once."""
        doc = Document(self.file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [LangchainDocument(page_content=text, metadata={"source": self.file_path})]

